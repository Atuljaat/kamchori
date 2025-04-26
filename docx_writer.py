from docx import Document
from docx.shared import Inches, RGBColor
from pdf_reader import extract_text_from_pdf
from getCode import main_response
from create_codeFile import create_code_file
from code_cleaner import split_algo_code_output, clean_output_for_C
from generate_output import output_to_image
import time
import re

# Initialize document
doc = Document()
doc.add_heading('COLLEGE KI FILE', 0)
doc.add_page_break()

# Extract questions from PDF
text = extract_text_from_pdf('PROGRAMS.pdf')
questions = re.split(r'(?<!\d)(?=\d{1,2}\.\s)', text)
questions = [q.strip() for q in questions if q.strip()]
text = questions

total_time = time.time()


def write_to_docx(filename):

    count = 1
    for index, line in enumerate(text):
        start_time = time.time()
        
        # Write Question
        heading = doc.add_heading(line, level=2)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
        
        code = (main_response(line, count).text)
        print(code)
        
        algo, code_string, output_string = split_algo_code_output(code)
        print("Printing output_text\n")
        print(output_string)

        # Write Algorithm
        heading = doc.add_heading('Algorithm', level=3)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
        code_para = doc.add_paragraph(algo)
        code_para.paragraph_format.space_after = 0
        code_para.paragraph_format.line_spacing = 1

        # Write Code
        heading = doc.add_heading('Code:', level=3)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
        code_para = doc.add_paragraph(code_string)
        code_para.paragraph_format.space_after = 0
        code_para.paragraph_format.line_spacing = 1

        # Write Output
        heading = doc.add_heading('Output:', level=3)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
        
        output_string = clean_output_for_C((str(output_string)).splitlines())
        image_path = output_to_image(output_string, f"images/output{index}.png")
        doc.add_picture(image_path)

        create_code_file(code_string, index)
        doc.add_paragraph("")
        doc.save(filename + '.docx')

        endtime = time.time()
        print(f"Time taken for question {index} : {endtime - start_time} seconds")
        
        # run_code(input, index)
        count += 1
